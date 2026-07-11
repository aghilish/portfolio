from __future__ import annotations

import re
from pathlib import Path

import yaml
from docx import Document
from docx.shared import Pt


def clean_markdown(text: str) -> str:
    """Convert basic markdown in YAML text to plain Word-friendly text."""
    if not isinstance(text, str):
        return str(text)

    # [label](url) -> label (url)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    # **bold** -> bold
    text = text.replace("**", "")
    # Keep readable output for escaped characters
    text = text.replace("\\(", "(").replace("\\)", ")")
    return text.strip()


def add_title_and_contact(doc: Document, cv_data: dict) -> None:
    title = doc.add_heading(cv_data.get("name", "CV"), level=0)
    if title.runs:
        title.runs[0].font.size = Pt(22)

    label = cv_data.get("label")
    if label:
        p = doc.add_paragraph(clean_markdown(label))
        if p.runs:
            p.runs[0].bold = True

    contact_parts: list[str] = []
    for key in ("location", "phone", "email", "website"):
        value = cv_data.get(key)
        if value:
            contact_parts.append(clean_markdown(value))
    if contact_parts:
        doc.add_paragraph(" | ".join(contact_parts))

    social = cv_data.get("social_networks", [])
    if social:
        social_parts = []
        for s in social:
            network = s.get("network")
            username = s.get("username")
            if network and username:
                social_parts.append(f"{clean_markdown(network)}: {clean_markdown(username)}")
        if social_parts:
            doc.add_paragraph(" | ".join(social_parts))


def add_section_entry(doc: Document, entry: dict) -> None:
    # One-line labeled entries (Skills/Languages)
    if "label" in entry and "details" in entry and "company" not in entry and "institution" not in entry:
        doc.add_paragraph(f"{clean_markdown(entry['label'])}: {clean_markdown(entry['details'])}")
        return

    heading_parts: list[str] = []
    if entry.get("position") and entry.get("company"):
        heading_parts.append(f"{clean_markdown(entry['position'])} - {clean_markdown(entry['company'])}")
    elif entry.get("institution"):
        degree = clean_markdown(entry.get("degree", ""))
        area = clean_markdown(entry.get("area", ""))
        tail = f"{degree} {area}".strip()
        heading_parts.append(f"{clean_markdown(entry['institution'])} - {tail}".strip(" -"))
    elif entry.get("name"):
        heading_parts.append(clean_markdown(entry["name"]))

    if heading_parts:
        p = doc.add_paragraph(" | ".join(heading_parts))
        if p.runs:
            p.runs[0].bold = True

    meta: list[str] = []
    if entry.get("location"):
        meta.append(clean_markdown(entry["location"]))
    if entry.get("start_date") or entry.get("end_date"):
        meta.append(f"{clean_markdown(entry.get('start_date', ''))} - {clean_markdown(entry.get('end_date', ''))}".strip())
    if meta:
        doc.add_paragraph(" | ".join([m for m in meta if m]))

    if entry.get("summary"):
        doc.add_paragraph(clean_markdown(entry["summary"]))

    highlights = entry.get("highlights", [])
    for h in highlights:
        doc.add_paragraph(clean_markdown(h), style="List Bullet")

    if entry.get("details") and "label" not in entry:
        doc.add_paragraph(clean_markdown(entry["details"]))


def add_section(doc: Document, section_name: str, section_content: list) -> None:
    doc.add_heading(clean_markdown(section_name), level=1)

    for item in section_content:
        if isinstance(item, str):
            doc.add_paragraph(clean_markdown(item), style="List Bullet")
            continue

        if isinstance(item, dict) and "bullet" in item:
            # Certifications and similar bullet sections
            doc.add_paragraph(clean_markdown(item["bullet"]), style="List Bullet")
            continue

        if isinstance(item, dict):
            add_section_entry(doc, item)
            doc.add_paragraph("")  # spacing between entries


def generate_docx(yaml_path: Path, output_path: Path) -> None:
    data = yaml.safe_load(yaml_path.read_text(encoding="utf-8"))
    cv_data = data.get("cv", {})
    sections = cv_data.get("sections", {})

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)

    add_title_and_contact(doc, cv_data)
    doc.add_paragraph("")

    for section_name, section_content in sections.items():
        if isinstance(section_content, list):
            add_section(doc, section_name, section_content)
        else:
            doc.add_heading(clean_markdown(section_name), level=1)
            doc.add_paragraph(clean_markdown(str(section_content)))

    doc.save(output_path)


def main() -> None:
    base = Path("/Users/saghili/personal/portfolio/cv")
    pairs = [
        (base / "Shahrooz_Aghili_CV_EN.yaml", base / "Shahrooz_Aghili_CV_EN.docx"),
        (base / "Shahrooz_Aghili_CV_DE.yaml", base / "Shahrooz_Aghili_CV_DE.docx"),
    ]
    for yaml_path, docx_path in pairs:
        generate_docx(yaml_path, docx_path)
        print(f"generated: {docx_path}")


if __name__ == "__main__":
    main()
